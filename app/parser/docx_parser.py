import docx


def extract_text_from_docx(file) -> str:
    """
    Reads text content from a DOCX file
    """
    doc = docx.Document(file)
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_content.append(cell_text)

    return " ".join(text_content)
